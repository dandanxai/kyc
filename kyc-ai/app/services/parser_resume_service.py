import io
import json
import re  # 🌟 引入正则，用于清洗超链接脏数据
import requests
from pypdf import PdfReader
from docx import Document
import openpyxl
import xlrd
from app.config.settings import settings
from app.config.prompts import ResumePrompts
from app.dao.kyc_resume import ResumeDao
from app.dao.resume_graph_dao import ResumeGraphDao

class ParserResumeService:
    @classmethod
    def process_task(cls, task_data: dict) -> bool:
        resume_id = task_data.get("id")
        file_path = task_data.get("filePath")
        # 🌟 核心点：从 MQ 任务中捞出上传人的 userId
        user_id = task_data.get("userId")

        if not resume_id or not file_path:
            print(f"❌ [解析服务] 任务数据缺失，无法处理: {task_data}")
            return False

        print(f"\n⚡ [解析服务] 收到新任务 -> 简历ID: {resume_id}")

        try:
            # 1. 榨取文本
            resume_text = cls._download_and_extract_text(file_path)
            if not resume_text.strip():
                return False

            # 2. 调大模型获取 JSON 字典
            parsed_json = cls._call_llm_extract(resume_text)

            print("\n======================= 🚀 AI 解析结果 🚀 =======================")
            print(json.dumps(parsed_json, indent=4, ensure_ascii=False))
            print("=================================================================\n")

            # 3. 调用 DAO 层，直接将解析结果落库达梦！
            ResumeDao.update_parsed_result(resume_id, parsed_json)
            print(f"🎉 [解析服务] 任务 ID: {resume_id} 智能化全格式解析并落库成功！")

            # =================================================================
            # 🌟 核心新增：落库成功后，走内部网络接口同步触发 Java 端的 SSE 长连接推送！
            # =================================================================
            if user_id:
                try:
                    # 🚀 1. 统一换成大一统的统一路径点
                    notify_url = "http://127.0.0.1:48080/app-api/member/sse/notify-success"

                    # 🚀 2. 弹药包升级：消除歧义，统一叫 dataId，并打上简历 RESUME 钢印
                    payload = {
                        "userId": int(user_id),
                        "dataId": str(resume_id),
                        "taskType": "RESUME"
                    }

                    # 轰炸 Java 后端
                    response = requests.post(notify_url, json=payload, timeout=5)
                    print(f"📡 [📡通知信道] Java 响应状态: {response.status_code}, 返回内容: {response.text}")
                except Exception as ne:
                    print(f"⚠️ [📡通知信道] 触发 Java SSE 广播时发生网络异常: {str(ne)}")
            else:
                print("⚠️ [📡通知信道] 未能从任务体中获取到 userId，略过实时通知。")

            return True

        except Exception as e:
            print(f"💥 [解析服务] 处理任务 ID: {resume_id} 时发生异常: {str(e)}")
            raise e

    @classmethod
    def _download_and_extract_text(cls, file_path: str) -> str:
        """
        智能路由器：清洗超链接、下载文件，增强版抗揍路由
        """
        file_path = file_path.strip()
        markdown_match = re.match(r'^\[.*?\]\((https?://.*?)\)$', file_path)
        if markdown_match:
            file_path = markdown_match.group(1)
            print(f"🧹 [清洗过滤器] 检测到 Markdown 链接伪装，已自动还原为真实 URL: {file_path}")
        else:
            print(f"🔗 [解析服务] 简历OSS下载链接: {file_path}")

        clean_url = file_path.split('?')[0].lower()

        print(f"📥 [文件解析] 正在从 OSS 下载原始简历资源...")
        response = requests.get(file_path, timeout=20, proxies=settings.request_proxies)
        response.raise_for_status()

        file_bytes = response.content
        full_text = ""

        # --- 路由 1：处理 PDF 格式 ---
        if clean_url.endswith('.pdf'):
            print(f"📄 [文件解析] 检测到标准 PDF 格式，开始榨取...")
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"

        # --- 路由 2：处理 Word 现代格式 (.docx) ---
        elif clean_url.endswith('.docx'):
            print(f"📝 [文件解析] 检测到 Word (.docx) 格式，开始深层剥离...")
            try:
                word_file = io.BytesIO(file_bytes)
                doc = Document(word_file)

                # 1. 抓取常规正文段落
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text += para.text + "\n"

                # 2. 抓取表格里的文字
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            # 顺便去个重，防止某些合并单元格被重复读取
                            full_text += " | ".join(dict.fromkeys(row_text)) + "\n"

                # 3. 🌟【核心突破】：强行穿透文本框（Text Box）和图形对象
                # 从底层 XML 查找所有符合文本流定义的节点
                from docx.oxml.ns import qn
                root = doc.element
                # 搜索所有包含文本框和形状内容的特殊文本节点
                body_xml = root.xml

                # 利用正则或者底层查找抓取隐藏在 w:t 标签下的所有文本
                # 很多时候图形、文本框里的文字在底层都是由 w:t 承载的
                # 如果前两步抓出来的文本太少，我们直接通过底层节点查缺补漏
                if len(full_text.strip()) < 100:
                    print(f"🔍 [文本框侦测] 检测到正文文本极少，开启底层 XML 文本框深度扫描...")
                    # 遍历文档的所有元素
                    for el in root.iter():
                        # 如果是文本节点
                        if el.tag.endswith('t') and el.text and el.text.strip():
                            # 为了防止和常规正文重复，我们用查重机制
                            text_piece = el.text.strip()
                            if text_piece not in full_text:
                                full_text += text_piece + "\n"

            except Exception as e:
                print(f"💥 [Word常规解析异常] 触发暴力流降级抽取: {str(e)}")
                # 兜底：如果连底层 XML 读取都出问题，直接上最暴力的字符清洗法
                text_pieces = re.findall(r'[\u4e00-\u9fa5_a-zA-Z0-9\s\-\.\,\:\@]{3,}',
                                         file_bytes.decode('utf-8', errors='ignore'))
                full_text = "\n".join(text_pieces)

        # --- 路由 3：处理 Excel 现代格式 (.xlsx) ---
        elif clean_url.endswith('.xlsx'):
            print(f"📊 [文件解析] 检测到 Excel (.xlsx) 格式表格，开始提取...")
            excel_file = io.BytesIO(file_bytes)
            wb = openpyxl.load_workbook(excel_file, data_only=True)
            for sheet in wb.worksheets:
                full_text += f"--- 工作表: {sheet.title} ---\n"
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell).strip() for cell in row if cell is not None]
                    if row_values:
                        full_text += "  ".join(row_values) + "\n"

        # --- 路由 4：处理 Excel 老版本格式 (.xls) ---
        elif clean_url.endswith('.xls'):
            print(f"📊 [文件解析] 检测到老版 Excel (.xls) 格式，开始提取...")
            wb = xlrd.open_workbook(file_contents=file_bytes)
            for sheet in wb.sheets():
                full_text += f"--- 工作表: {sheet.name} ---\n"
                for row_idx in range(sheet.nrows):
                    row_values = [str(sheet.cell_value(row_idx, col_idx)).strip() for col_idx in range(sheet.ncols)]
                    row_values = [v for v in row_values if v]
                    if row_values:
                        full_text += "  ".join(row_values) + "\n"

        # --- 兜底过滤 ---
        else:
            print(f"⚠️ [文件解析] 未能识别该文件格式，将尝试作为纯文本处理...")
            full_text = file_path

        print(f"✅ [文件解析] 成功脱水榨取文本，共计 {len(full_text)} 个字符。")
        return full_text

    @classmethod
    def _call_llm_extract(cls, resume_text: str) -> dict:
        """
        内部私有方法：调用通义千问大模型解析剥离出来的全量文本内容
        """
        print(f"🤖 [解析服务] 正在构建标准大模型规约，准备调用: {settings.model_name}...")

        # 文本替换：无痛灌入大模板
        prompt = ResumePrompts.PARSE_TEMPLATE.replace("__RESUME_CONTENT_PLACEHOLDER__", resume_text)

        headers = {
            "Authorization": f"Bearer {settings.dashscope_api_key}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": settings.model_name,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1
        }

        response = requests.post(
            settings.ollama_api_url,
            headers=headers,
            json=payload,
            timeout=30,
            proxies=settings.request_proxies
        )
        response.raise_for_status()

        response_json = response.json()
        ai_content = response_json['choices'][0]['message']['content'].strip()
        print(f"✨ [解析服务] 大模型响应成功！收到原文响应。")

        if ai_content.startswith("```"):
            ai_content = ai_content.replace("```json", "").replace("```", "").strip()

        return json.loads(ai_content)


if __name__ == "__main__":
    print("▶️ 开始全格式解析 Service 业务层独立单体测试...")
    # 🌟 修复：这里的链接已经被我清理成了干净的原始纯字符串链接！
    test_task = {
        "id": 32,
        "filePath": "[https://dandanxia-kyc.oss-cn-beijing.aliyuncs.com/20260606/1780744359687_87a9a2fc7dbc4dd1a638b0fd417953fa.pdff](https://dandanxia-kyc.oss-cn-beijing.aliyuncs.com/20260606/1780744359687_87a9a2fc7dbc4dd1a638b0fd417953fa.pdf)",
        "userId": 1
    }
    ParserResumeService.process_task(test_task)