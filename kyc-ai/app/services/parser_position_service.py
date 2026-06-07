import io
import json
import re  # 🌟 引入正则，用于清洗超链接脏数据
import requests
from pypdf import PdfReader
from docx import Document
import openpyxl
import xlrd
from app.config.settings import settings
from app.config.prompts import JobPrompts  # 🌟 切换为岗位提示词模板
from app.dao.kyc_position import PositionDao  # 🌟 引入刚刚写好的岗位 Dao 层


class ParserPositionService:
    @classmethod
    def process_task(cls, task_data: dict) -> bool:
        position_id = task_data.get("id")
        file_path = task_data.get("filePath")
        # 🌟 从 MQ 任务中捞出上传该职位的 HR/企业用户 userId (用于分发 SSE 报喜信号)
        user_id = task_data.get("userId")

        if not position_id or not file_path:
            print(f"❌ [岗位解析] 任务数据缺失，无法处理: {task_data}")
            return False

        print(f"\n⚡ [岗位解析] 收到新招聘文档任务 -> 岗位ID: {position_id}")

        try:
            # 1. 压榨出招聘简章、Word描述里的纯文本内容 (100% 沿用无损路由器)
            job_text = cls._download_and_extract_text(file_path)
            if not job_text.strip():
                return False

            # 2. 调用通义千问/DeepSeek大模型按照 B 端 JSON 骨架抽取画像
            parsed_json = cls._call_llm_extract(job_text)

            print("\n======================= 🚀 AI 岗位解构结果 🚀 =======================")
            print(json.dumps(parsed_json, indent=4, ensure_ascii=False))
            print("=================================================================\n")

            # 3. 🎯【核心归位】：调用 Dao 层，将结构化标签及全量 JSON 砸进达梦数据库！
            PositionDao.update_parsed_result(position_id, parsed_json)
            print(f"🎉 [岗位解析] 任务 ID: {position_id} 达梦库硬性指标同步成功！")

            # =================================================================
            # 4. 核心对齐：岗位全脱水落库成功后，跨网络触发 Java 端 B 端的 SSE 推送！
            # =================================================================
            if user_id:
                try:
                    # 🚀 1. 同样换成大一统的统一路径点
                    notify_url = "http://127.0.0.1:48080/app-api/member/sse/notify-success"

                    # 🚀 2. 弹药包升级：统一叫 dataId，并打上岗位 POSITION 钢印
                    payload = {
                        "userId": int(user_id),
                        "dataId": str(position_id),
                        "taskType": "POSITION"
                    }

                    # 轰炸 Java 后端
                    response = requests.post(notify_url, json=payload, timeout=5)
                    print(f"📡 [📡 B端通知信道] Java 响应状态: {response.status_code}, 返回内容: {response.text}")
                except Exception as ne:
                    print(f"⚠️ [📡 B端通知信道] 触发 Java B端 SSE 广播时网络瘫痪: {str(ne)}")
            else:
                print("⚠️ [📡 B端通知信道] 未能从任务体中获取到 userId，跳过 SSE 实时通报。")

            return True

        except Exception as e:
            print(f"💥 [岗位解析] 处理任务 ID: {position_id} 时发生未知暴击: {str(e)}")
            raise e

    @classmethod
    def _download_and_extract_text(cls, file_path: str) -> str:
        """
        100% 像素级对齐的智能文件路由器：清洗超链接、多格式压榨抽取
        """
        file_path = file_path.strip()
        markdown_match = re.match(r'^\[.*?\]\((https?://.*?)\)$', file_path)
        if markdown_match:
            file_path = markdown_match.group(1)
            print(f"🧹 [清洗过滤器] 检测到 Markdown 链接伪装，已自动还原为真实 URL: {file_path}")
        else:
            print(f"🔗 [岗位解析] 岗位描述OSS下载链接: {file_path}")

        clean_url = file_path.split('?')[0].lower()

        print(f"📥 [文件解析] 正在从 OSS 下载原始岗位 JD 资源...")
        response = requests.get(file_path, timeout=20, proxies=settings.request_proxies)
        response.raise_for_status()

        file_bytes = response.content
        full_text = ""

        # --- 路由 1：处理 PDF 招聘需求 ---
        if clean_url.endswith('.pdf'):
            print(f"📄 [文件解析] 检测到 PDF 格式招聘需求，开始榨取...")
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"

        # --- 路由 2：处理 Word 招聘需求 (.docx) ---
        elif clean_url.endswith('.docx'):
            print(f"📝 [文件解析] 检测到 Word (.docx) 格式需求文档，开始深层剥离...")
            try:
                word_file = io.BytesIO(file_bytes)
                doc = Document(word_file)

                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text += para.text + "\n"

                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            full_text += " | ".join(dict.fromkeys(row_text)) + "\n"

                if len(full_text.strip()) < 100:
                    print(f"🔍 [文本框侦测] 检测到招聘正文文本极少，开启底层 XML 深度扫描...")
                    root = doc.element
                    for el in root.iter():
                        if el.tag.endswith('t') and el.text and el.text.strip():
                            text_piece = el.text.strip()
                            if text_piece not in full_text:
                                full_text += text_piece + "\n"

            except Exception as e:
                print(f"💥 [Word常规岗位解析异常] 触发暴力流降级抽取: {str(e)}")
                text_pieces = re.findall(r'[\u4e00-\u9fa5_a-zA-Z0-9\s\-\.\,\:\@]{3,}',
                                         file_bytes.decode('utf-8', errors='ignore'))
                full_text = "\n".join(text_pieces)

        # --- 路由 3：处理 Excel 批量发布格式 (.xlsx) ---
        elif clean_url.endswith('.xlsx'):
            print(f"📊 [文件解析] 检测到 Excel (.xlsx) 岗位大表，开始提取...")
            excel_file = io.BytesIO(file_bytes)
            wb = openpyxl.load_workbook(excel_file, data_only=True)
            for sheet in wb.worksheets:
                full_text += f"--- 工作表: {sheet.title} ---\n"
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell).strip() for cell in row if cell is not None]
                    if row_values:
                        full_text += "  ".join(row_values) + "\n"

        # --- 路由 4：处理 Excel 老版本格式 (.xls) ---
        elif clean_url.endswith('.xls'):
            print(f"📊 [文件解析] 检测到老版 Excel (.xls) 岗位表，开始提取...")
            wb = xlrd.open_workbook(file_contents=file_bytes)
            for sheet in wb.sheets():
                full_text += f"--- 工作表: {sheet.name} ---\n"
                for row_idx in range(sheet.nrows):
                    row_values = [str(sheet.cell_value(row_idx, col_idx)).strip() for col_idx in range(sheet.ncols)]
                    row_values = [v for v in row_values if v]
                    if row_values:
                        full_text += "  ".join(row_values) + "\n"

        # --- 兜底过滤 ---
        else:
            print(f"⚠️ [文件解析] 未能识别该文件格式，直接作为文本路由兜底...")
            full_text = file_path

        print(f"✅ [文件解析] 成功脱水榨取岗位文本，共计 {len(full_text)} 个字符。")
        return full_text

    @classmethod
    def _call_llm_extract(cls, job_text: str) -> dict:
        """
        内部私有方法：调用大模型解析剥离出来的全量岗位文本内容
        """
        print(f"🤖 [岗位解析] 正在构建标准岗位大模型规约，准备调用: {settings.model_name}...")

        prompt = JobPrompts.ANALYZE_TEMPLATE.replace("__JOB_CONTENT_PLACEHOLDER__", job_text)

        headers = {
            "Authorization": f"Bearer {settings.dashscope_api_key}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": settings.model_name,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1
        }

        response = requests.post(
            settings.ollama_api_url,
            headers=headers,
            json=payload,
            timeout=30,
            proxies=settings.request_proxies
        )
        response.raise_for_status()

        response_json = response.json()
        ai_content = response_json['choices'][0]['message']['content'].strip()
        print(f"✨ [岗位解析] 大模型岗位解构大成功！")

        if ai_content.startswith("```"):
            ai_content = ai_content.replace("```json", "").replace("```", "").strip()

        return json.loads(ai_content)


if __name__ == "__main__":
    print("▶️ 开始岗位招聘解析 Service 业务层独立单体调试测试...")
    test_job_task = {
        "id": 88,
        "filePath": "[https://dandanxia-kyc.oss-cn-beijing.aliyuncs.com/20260606/Java开发工程师岗位需求.docx](https://dandanxia-kyc.oss-cn-beijing.aliyuncs.com/20260606/Java开发工程师岗位需求.docx)",
        "userId": 3
    }
    ParserJobService.process_task(test_job_task)